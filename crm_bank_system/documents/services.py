from django.contrib.auth import get_user_model
from django.db.models import Count
from docx import Document
from .dto import ReportDTO
from .models import Report
from django.conf import settings

User = get_user_model()

def create_report_document(report):
    document = Document()

    document.add_heading('Жалоба', 0)

    document.add_paragraph(f'Пользователь: {report.user.username}')
    document.add_paragraph(f'Описание проблемы: {report.description}')
    document.add_paragraph(f'URL скриншота: {report.screenshot}')
    document.add_paragraph(f'Статус жалобы: {report.get_status_display()}')
    document.add_paragraph(f'Дата создания: {report.created_at}')

    document_path = f'media/reports/{report.id}.docx'
    document.save(document_path)

    report.document_url = f'{settings.MEDIA_URL}reports/{report.id}.docx'
    report.save()

    return document_path

def create_report(user, description, screenshot):
    user = User.objects.get(id=user)
    document_dto = ReportDTO(user=user,
                            description=description,
                            screenshot=screenshot)
    created_report = Report.objects.create_report(document_dto)
    return created_report

def reports_list():
    users = User.objects.annotate(report_count=Count('reports')).prefetch_related('reports').all()
    user_reports = []

    for user in users:
        reports = user.reports.all()
        if reports.exists():
            user_reports.append({
                'username': user.username,
                'reports': reports
            }
            )
    return user_reports


def get_detail_report(report_id):
    report = Report.objects.get_detail_report(pk=report_id)

    if not report:
        raise ValueError(f'Не существует такой жалобы')

    return report